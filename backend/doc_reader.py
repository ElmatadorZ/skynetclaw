"""
doc_reader.py — OX-DOC-UPLOAD-1
===============================
Extract readable text from uploaded documents / images so SkynetClaw can read
and use them. Pure-python extractors (pypdf / python-docx / openpyxl) + native
text + HTML (bs4). Images: returns metadata; text-from-image (OCR) needs an OCR
engine or a vision model (not available with the text-only local model).

License: Apache-2.0 — ElmatadorZ / THE HOUSE
"""
from __future__ import annotations

import csv
import io
import json
import os
from typing import Any, Dict

TEXT_EXT = {".txt", ".md", ".markdown", ".log", ".json", ".csv", ".tsv",
            ".py", ".js", ".ts", ".html", ".htm", ".css", ".xml", ".yaml",
            ".yml", ".ini", ".cfg", ".sql", ".sh", ".bat", ".ps1"}
IMG_EXT = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff"}
MAX_CHARS = 200_000   # safety cap so a huge doc can't blow the prompt/budget


def _pdf(path: str) -> str:
    from pypdf import PdfReader
    r = PdfReader(path)
    out = []
    for i, pg in enumerate(r.pages, 1):
        try:
            t = pg.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            out.append(f"--- page {i} ---\n{t.strip()}")
    return "\n\n".join(out) or "[PDF has no extractable text — likely a scanned image PDF; OCR needed]"


def _docx(path: str) -> str:
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts) or "[DOCX has no text]"


def _xlsx(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"=== sheet: {ws.title} ===")
        for row in ws.iter_rows(values_only=True):
            vals = ["" if v is None else str(v) for v in row]
            if any(vals):
                out.append(" | ".join(vals))
    return "\n".join(out) or "[XLSX empty]"


def _html(path: str) -> str:
    from bs4 import BeautifulSoup
    raw = open(path, encoding="utf-8", errors="replace").read()
    soup = BeautifulSoup(raw, "html.parser")
    for t in soup(["script", "style"]):
        t.decompose()
    return "\n".join(line for line in (l.strip() for l in soup.get_text("\n").splitlines()) if line)


def _find_tesseract() -> str:
    """Locate the tesseract binary. PATH first — that is how it is installed on
    Linux (apt/dnf/pacman), macOS (brew), and increasingly on Windows too.
    The explicit lists are fallbacks for installers that skip PATH."""
    import shutil
    cmd = shutil.which("tesseract")
    if cmd:
        return cmd
    if os.name == "nt":
        cands = (r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                 r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe")
    else:
        cands = ("/usr/bin/tesseract",            # Debian/Ubuntu, Fedora, Arch
                 "/usr/local/bin/tesseract",      # source builds, Intel brew
                 "/opt/homebrew/bin/tesseract",   # Apple-silicon brew
                 "/snap/bin/tesseract")
    for p in cands:
        if os.path.exists(p):
            return p
    return ""


def _find_tessdata() -> str:
    """A tessdata dir that contains Thai (tha.traineddata)."""
    home = os.path.expanduser("~")
    cands = [os.environ.get("TESSDATA_PREFIX", ""),
             os.path.join(home, "llamacpp_test", "tessdata")]
    if os.name == "nt":
        cands += [r"C:\Program Files\Tesseract-OCR\tessdata",
                  r"C:\Program Files (x86)\Tesseract-OCR\tessdata"]
    else:
        cands += ["/usr/share/tesseract-ocr/5/tessdata",   # Debian/Ubuntu (v5)
                  "/usr/share/tesseract-ocr/4.00/tessdata",  # older Debian
                  "/usr/share/tessdata",                     # Fedora/Arch
                  "/usr/local/share/tessdata",
                  "/opt/homebrew/share/tessdata"]            # Apple-silicon brew
    for d in cands:
        if d and os.path.exists(os.path.join(d, "tha.traineddata")):
            return d
    return ""


def _image(path: str) -> str:
    info = {}
    try:
        from PIL import Image
        with Image.open(path) as im:
            info = {"format": im.format, "size": list(im.size), "mode": im.mode}
    except Exception:
        pass
    cmd = _find_tesseract()
    if not cmd:
        return (f"[image uploaded — {info}. อ่านข้อความในรูปไม่ได้: ไม่พบ Tesseract OCR. "
                f"ติดตั้ง Tesseract หรือใช้ vision model.]")
    try:
        import pytesseract
        from PIL import Image
        pytesseract.pytesseract.tesseract_cmd = cmd
        tdir = _find_tessdata()
        if tdir:
            os.environ["TESSDATA_PREFIX"] = tdir   # avoids --tessdata-dir quoting issues
        lang = "tha+eng" if tdir else "eng"
        txt = pytesseract.image_to_string(Image.open(path), lang=lang).strip()
        if txt:
            return f"[image OCR · {lang} · {info}]\n{txt}"
        return f"[image {info} — OCR ทำงานแล้วแต่ไม่พบข้อความในรูป (อาจเป็นภาพถ่าย/กราฟ ไม่ใช่ตัวอักษร)]"
    except Exception as e:
        return f"[image uploaded — {info}. OCR error: {str(e)[:140]}]"


def extract_text(path: str) -> Dict[str, Any]:
    """Return {ok, kind, chars, text, truncated, error}."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            text, kind = _pdf(path), "pdf"
        elif ext == ".docx":
            text, kind = _docx(path), "docx"
        elif ext in (".xlsx", ".xlsm"):
            text, kind = _xlsx(path), "xlsx"
        elif ext in (".html", ".htm"):
            text, kind = _html(path), "html"
        elif ext in IMG_EXT:
            text, kind = _image(path), "image"
        elif ext in TEXT_EXT or ext == "":
            text, kind = open(path, encoding="utf-8", errors="replace").read(), "text"
        elif ext == ".doc":
            return {"ok": False, "kind": "doc", "error": "legacy .doc not supported — save as .docx or .pdf"}
        else:
            # try as text; if binary, refuse
            try:
                text, kind = open(path, encoding="utf-8").read(), "text"
            except Exception:
                return {"ok": False, "kind": ext or "binary",
                        "error": f"unsupported file type '{ext}'"}
        truncated = len(text) > MAX_CHARS
        if truncated:
            text = text[:MAX_CHARS] + f"\n…[truncated, {kind} longer than {MAX_CHARS} chars]"
        return {"ok": True, "kind": kind, "chars": len(text), "text": text, "truncated": truncated}
    except Exception as e:
        return {"ok": False, "kind": ext, "error": f"{type(e).__name__}: {str(e)[:200]}"}
